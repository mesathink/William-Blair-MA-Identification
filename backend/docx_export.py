"""
Builds a formatted Word doc of all delivered acquirer rationales from a completed report run.
Pure function of (target_profile, report_run) -- no FastAPI/network/DataFrame dependency, so it's
testable on its own with fabricated data, same as report_agent.py's other data-shaping functions.
"""

from __future__ import annotations
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from data_processing import FIELD_LABELS

CONVICTION_COLORS = {
    "High": RGBColor(0x1F, 0x8A, 0x5F),
    "Medium": RGBColor(0xB8, 0x86, 0x0B),
    "Low": RGBColor(0xC0, 0x39, 0x2B),
}
# data_processing.FIELD_LABELS covers the numeric fields; geography/ownership aren't part of that
# form-field set (no bands), so they're added here where the target profile is walked field-by-field.
OPTIONAL_FIELD_LABELS = [
    (f, FIELD_LABELS[f]) for f in ("target_revenue_mm", "target_ebitda_mm", "ebitda_margin_pct", "revenue_growth_pct")
] + [("geography", "Geography"), ("target_ownership_pre", "Ownership")]

def _add_section(doc: Document, title: str, body: str) -> None:
    doc.add_heading(title, level=2)
    doc.add_paragraph(body)

# One title page (target profile + delivered/requested summary) followed by one page per acquirer, in the same stack-ranked order the UI shows them
def build_report_docx(target_profile: dict, report_run: dict) -> io.BytesIO:
    doc = Document()

    title = doc.add_heading("Acquirer Identification Rationales", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    deal_size_value = target_profile["deal_size_mm"]["value"]
    deal_size_str = f"${deal_size_value / 1000:g}B" if deal_size_value >= 1000 else f"${deal_size_value:g}M"

    target_line = doc.add_paragraph()
    target_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target_line.add_run(f"Target: {target_profile['sector']} · Deal Size {deal_size_str}").bold = True

    optional_bits = [
        f"{label}: {v.get('label', v.get('value'))}"
        for field, label in OPTIONAL_FIELD_LABELS
        if (v := target_profile.get(field))
    ]
    if optional_bits:
        doc.add_paragraph(" · ".join(optional_bits)).alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary.add_run(f"{report_run['delivered_count']} of {report_run['requested_count']} acquirers")
    summary_run.italic = True

    for entry in report_run["delivered"]:
        doc.add_page_break()
        r = entry["report"]

        doc.add_heading(entry["acquirer"], level=1)

        subtitle = doc.add_paragraph()
        subtitle.add_run(f"Tier {entry['tier']}").bold = True
        subtitle.add_run("   ·   Conviction: ")
        conviction_run = subtitle.add_run(r["conviction_level"].upper())
        conviction_run.bold = True
        conviction_run.font.color.rgb = CONVICTION_COLORS.get(r["conviction_level"], RGBColor(0, 0, 0))

        scores = doc.add_paragraph()
        secondary = entry["secondary_score"]
        scores_run = scores.add_run(
            f"Primary {entry['primary_score']}%  ·  Secondary {secondary if secondary is not None else '—'}%  ·  "
            f"Strategic {entry['strategic_score']}%  ·  Relevancy {entry['relevancy_score']}%"
        )
        scores_run.font.size = Pt(9)
        scores_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        _add_section(doc, "Acquirer Overview", r["acquirer_overview"])
        _add_section(doc, "Strategic Fit Thesis", r["strategic_fit_thesis"])
        _add_section(doc, "Precedent Activity", r["precedent_activity"])
        _add_section(doc, "Valuation Context", r["valuation_context"])

        doc.add_heading("Risk Flags", level=2)
        for flag in r["risk_flags"]:
            doc.add_paragraph(flag, style="List Bullet")

        _add_section(doc, "Conviction Rationale", r["conviction_rationale"])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
